import docx

def generate_tailored_resume(missing_skills, output_filename="tailored_resume.docx"):
    """Creates a new Word document resume injecting missing keywords."""
    
    # 1. Create a blank Word document in memory
    doc = docx.Document()
    
    # 2. Add Header / Contact Info
    doc.add_heading('David - Software Developer', level=0)
    
    # 3. Add a Professional Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        "Highly motivated developer with a strong foundation in programming. "
        "Continuously expanding skill set to build efficient, scalable, and automated applications."
    )
    
    # 4. Add Skills Section
    doc.add_heading('Technical Skills', level=1)
    
    # We add the skills it already found...
    doc.add_paragraph("Core Competencies: Git, Data Analysis")
    
    # ...and inject the missing skills to beat the HR filter!
    if missing_skills:
        # Format the list nicely with capital letters
        formatted_skills = [skill.title() for skill in missing_skills]
        doc.add_paragraph(f"Recently Acquired Skills: {', '.join(formatted_skills)}")
        
    # 5. Add a placeholder for Experience
    doc.add_heading('Experience', level=1)
    doc.add_paragraph("Software Developer - Automated Job Agent Project (2026)")
    
    # 6. Save the physical file to your hard drive
    doc.save(output_filename)
    
# ==========================================
# MAIN SCRIPT EXECUTION
# ==========================================
if __name__ == "__main__":
    print("Initializing Stage 6: Resume Generator...")
    
    # These are the exact skills your analyzer said were missing!
    missing_from_analyzer = ['api', 'python', 'requests', 'json', 'github']
    
    # Run the engine
    generate_tailored_resume(missing_from_analyzer)
    
    print("Success! A new tailored resume has been generated to beat the HR filters.")